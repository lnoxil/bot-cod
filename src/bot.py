import asyncio
import json
import logging
import os
import re
import zipfile
from datetime import datetime, timezone
from dataclasses import asdict, dataclass, field, fields, replace
from pathlib import Path

import discord
from aiohttp import ClientSession, web
from discord import ButtonStyle
from discord.ext import commands
from dotenv import load_dotenv
from telegram import InlineKeyboardButton, InlineKeyboardMarkup, Update, BotCommand
from telegram.ext import (
    Application,
    CallbackQueryHandler,
    CommandHandler,
    ContextTypes,
    MessageHandler,
    filters,
)

load_dotenv()
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
)
logger = logging.getLogger("bridge")


def env_int(name: str, default: int | None = None) -> int | None:
    raw = os.getenv(name)
    if raw is None or raw == "":
        return default
    return int(raw)


def env_int_list(name: str) -> set[int]:
    raw = os.getenv(name, "").strip()
    if not raw:
        return set()
    return {int(x.strip()) for x in raw.split(",") if x.strip()}


def sanitize_filename(name: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9._-]", "_", name).strip("._")
    return cleaned or "file"

DISCORD_TOKEN = os.getenv("DISCORD_TOKEN")
TELEGRAM_TOKEN = os.getenv("TELEGRAM_TOKEN")
SUPPORT_ROLE_ID = env_int("SUPPORT_ROLE_ID")
TICKET_CATEGORY_ID = env_int("TICKET_CATEGORY_ID")
ORDER_PREFIX = os.getenv("ORDER_CHANNEL_PREFIX", "order")
SUPPORT_PREFIX = os.getenv("SUPPORT_CHANNEL_PREFIX", "support")
STATE_FILE = Path(os.getenv("STATE_FILE", "state/tickets.json"))
POSTS_FILE = Path(os.getenv("POSTS_FILE", "state/posts.json"))
USER_LINKS_FILE = Path(os.getenv("USER_LINKS_FILE", "state/user_links.json"))
TG_ROLES_FILE = Path(os.getenv("TG_ROLES_FILE", "state/tg_roles.json"))
TG_ADMIN_IDS = env_int_list("TG_ADMIN_IDS")
WEB_HOST = os.getenv("WEB_HOST", "0.0.0.0")
WEB_PORT = env_int("WEB_PORT", 8080) or 8080
EDITOR_HTML_PATH = Path(os.getenv("EDITOR_HTML_PATH", "web/editor.html"))
ARCHIVE_DIR = Path(os.getenv("ARCHIVE_DIR", "state/archives"))


@dataclass
class ExtraBlock:
    title: str = ""
    description: str = ""
    color_hex: str = "2ECC71"
    image_url: str | None = None
    image_position: str = "bottom"


@dataclass
class PanelButton:
    label: str = "Button"
    emoji: str = ""
    style: str = "secondary"
    action: str = "none"  # order | support | url | none
    url: str | None = None
    row: int = 0


@dataclass
class TicketBinding:
    ticket_channel_id: int
    opener_discord_id: int
    ticket_type: str
    digest_message_ids: dict[str, int] = field(default_factory=dict)


@dataclass
class SavedPost:
    name: str
    channel_id: int
    title: str
    description: str
    color_hex: str = "2ECC71"
    image_url: str | None = None
    image_position: str = "bottom"

    last_message_id: int | None = None
    is_ticket_panel: bool = False

    order_label: str = "ORDER"
    order_emoji: str = "🧾"
    order_style: str = "success"

    support_label: str = "SUPPORT"
    support_emoji: str = "🛟"
    support_style: str = "primary"

    auto_gradient: bool = False
    gradient_start_hex: str = "2ECC71"
    gradient_end_hex: str = "5865F2"

    auto_order_message: str = "Спасибо за заказ! Опишите задачу и бюджет."
    auto_support_message: str = "Спасибо за обращение в саппорт! Опишите проблему подробно."

    panel_buttons: list[PanelButton] = field(default_factory=list)
    extra_blocks: list[ExtraBlock] = field(default_factory=list)


class JsonStore:
    def __init__(self, path: Path):
        self.path = path
        self.path.parent.mkdir(parents=True, exist_ok=True)

    def _load(self) -> dict:
        if not self.path.exists():
            return {}
        return json.loads(self.path.read_text(encoding="utf-8"))

    def _save(self, payload: dict) -> None:
        self.path.write_text(
            json.dumps(payload, indent=2, ensure_ascii=False),
            encoding="utf-8",
        )


class TicketStore(JsonStore):
    def __init__(self, path: Path):
        super().__init__(path)
        self._bindings: dict[int, TicketBinding] = {}
        self.load()

    def load(self) -> None:
        data = self._load()
        self._bindings = {int(k): TicketBinding(**v) for k, v in data.items()}

    def save(self) -> None:
        self._save({k: asdict(v) for k, v in self._bindings.items()})

    def set(self, item: TicketBinding) -> None:
        self._bindings[item.ticket_channel_id] = item
        self.save()

    def get(self, channel_id: int) -> TicketBinding | None:
        return self._bindings.get(channel_id)

    def remove(self, channel_id: int) -> None:
        self._bindings.pop(channel_id, None)
        self.save()




def normalize_saved_post(raw: dict) -> SavedPost:
    allowed = {f.name for f in fields(SavedPost)}
    data = dict(raw)

    # Legacy migration: old split_* fields -> extra block continuation
    split_enabled = bool(data.pop("split_enabled", False))
    split_title = str(data.pop("split_title", "")).strip()
    split_description = str(data.pop("split_description", "")).strip()
    split_color_hex = str(data.pop("split_color_hex", data.get("color_hex", "2ECC71"))).strip()

    raw_blocks = data.get("extra_blocks", []) or []
    blocks: list[ExtraBlock] = []
    for block in raw_blocks:
        if isinstance(block, dict):
            block_allowed = {f.name for f in fields(ExtraBlock)}
            clean_block = {k: v for k, v in block.items() if k in block_allowed}
            blocks.append(ExtraBlock(**clean_block))

    raw_panel_buttons = data.get("panel_buttons", []) or []
    panel_buttons: list[PanelButton] = []
    for btn in raw_panel_buttons:
        if isinstance(btn, dict):
            try:
                row = int(btn.get("row", 0))
            except Exception:
                row = 0
            action = str(btn.get("action", "none")).strip().lower()
            if action == "link":
                action = "url"
            if action not in {"order", "support", "url", "none"}:
                action = "none"
            panel_buttons.append(
                PanelButton(
                    label=str(btn.get("label", "Button")),
                    emoji=str(btn.get("emoji", "")).strip(),
                    style=normalize_style_name(str(btn.get("style", "secondary"))),
                    action=action,
                    url=str(btn.get("url", "")).strip() or None,
                    row=max(0, min(4, row)),
                )
            )

    if split_enabled and (split_title or split_description):
        blocks.append(
            ExtraBlock(
                title=split_title,
                description=split_description,
                color_hex=split_color_hex or "2ECC71",
                image_url=None,
                image_position="bottom",
            )
        )

    data["panel_buttons"] = panel_buttons
    data["extra_blocks"] = blocks
    data["order_style"] = normalize_style_name(str(data.get("order_style", "success")))
    data["support_style"] = normalize_style_name(str(data.get("support_style", "primary")))
    data["order_emoji"] = str(data.get("order_emoji", "🧾")).strip()
    data["support_emoji"] = str(data.get("support_emoji", "🛟")).strip()
    clean = {k: v for k, v in data.items() if k in allowed}

    # Minimal required fallback for broken legacy records
    clean.setdefault("name", str(raw.get("name", "post")))
    clean.setdefault("channel_id", int(raw.get("channel_id", 0)))
    clean.setdefault("title", str(raw.get("title", "")))
    clean.setdefault("description", str(raw.get("description", "")))

    return SavedPost(**clean)

class PostStore(JsonStore):
    def __init__(self, path: Path):
        super().__init__(path)
        self._posts: dict[str, SavedPost] = {}
        self.load()

    def load(self) -> None:
        data = self._load()
        out: dict[str, SavedPost] = {}

        if isinstance(data, list):
            entries = []
            for i, item in enumerate(data):
                if isinstance(item, dict):
                    nm = str(item.get("name", f"post_{i}")).strip() or f"post_{i}"
                    entries.append((nm, item))
        elif isinstance(data, dict):
            entries = list(data.items())
        else:
            entries = []

        for name, raw in entries:
            if not isinstance(raw, dict):
                continue
            merged = dict(raw)
            merged.setdefault("name", str(name))
            try:
                post = normalize_saved_post(merged)
                out[post.name] = post
            except Exception:
                logger.exception("Skip invalid saved post '%s' during load", name)
        self._posts = out

    def save(self) -> None:
        payload = {name: asdict(post) for name, post in self._posts.items()}
        self._save(payload)

    def set(self, post: SavedPost) -> None:
        self._posts[post.name] = post
        self.save()

    def get(self, name: str) -> SavedPost | None:
        return self._posts.get(name)

    def list_posts(self) -> list[SavedPost]:
        return [self._posts[name] for name in sorted(self._posts)]


class UserLinkStore(JsonStore):
    def __init__(self, path: Path):
        super().__init__(path)
        self.discord_to_tg_chat: dict[int, int] = {}
        self.load()

    def load(self) -> None:
        data = self._load()
        self.discord_to_tg_chat = {
            int(k): int(v) for k, v in data.get("discord_to_tg_chat", {}).items()
        }

    def save(self) -> None:
        self._save({"discord_to_tg_chat": self.discord_to_tg_chat})

    def link(self, discord_user_id: int, tg_chat_id: int) -> None:
        self.discord_to_tg_chat[discord_user_id] = tg_chat_id
        self.save()

    def get_tg_chat(self, discord_user_id: int) -> int | None:
        return self.discord_to_tg_chat.get(discord_user_id)


class TgRoleStore(JsonStore):
    def __init__(self, path: Path):
        super().__init__(path)
        self.roles: dict[int, str] = {}
        self.role_chats: dict[int, int] = {}
        self.load()

    def load(self) -> None:
        data = self._load()
        self.roles = {int(k): str(v) for k, v in data.get("roles", {}).items()}
        self.role_chats = {
            int(k): int(v) for k, v in data.get("role_chats", {}).items()
        }

    def save(self) -> None:
        self._save({"roles": self.roles, "role_chats": self.role_chats})

    def set_role(self, tg_user_id: int, role: str, chat_id: int | None = None) -> None:
        self.roles[tg_user_id] = role
        if chat_id is not None:
            self.role_chats[tg_user_id] = chat_id
        self.save()

    def register_chat(self, tg_user_id: int, chat_id: int) -> None:
        self.role_chats[tg_user_id] = chat_id
        self.save()

    def get_role(self, tg_user_id: int) -> str:
        return self.roles.get(tg_user_id, "viewer")

    def users_with_roles(self, allowed: set[str]) -> list[int]:
        return [uid for uid, role in self.roles.items() if role in allowed]

    def chats_with_roles(self, allowed: set[str]) -> set[int]:
        out: set[int] = set()
        for uid, role in self.roles.items():
            if role in allowed and uid in self.role_chats:
                out.add(self.role_chats[uid])
        return out


class CloseTicketButton(discord.ui.View):
    def __init__(self, bot: "BridgeBot"):
        super().__init__(timeout=None)
        self.bot = bot

    @discord.ui.button(
        label="Close ticket",
        style=ButtonStyle.danger,
        custom_id="ticket_close",
        emoji="🗑️",
    )
    async def close_ticket(self, interaction: discord.Interaction, _: discord.ui.Button) -> None:
        if not interaction.channel:
            return
        binding = self.bot.ticket_store.get(interaction.channel.id)
        if binding is None:
            await interaction.response.send_message("Это не тикет-канал.", ephemeral=True)
            return

        await interaction.response.send_message("Тикет закрывается…", ephemeral=True)
        await self.bot.close_ticket(interaction.channel, binding, closed_by=interaction.user)


SUPPORTED_STYLES = {"primary", "secondary", "success", "danger"}


def normalize_style_name(name: str) -> str:
    nm = str(name or "").strip().lower()
    return nm if nm in SUPPORTED_STYLES else "secondary"


def style_from_name(name: str) -> ButtonStyle:
    mapping = {
        "primary": ButtonStyle.primary,
        "secondary": ButtonStyle.secondary,
        "success": ButtonStyle.success,
        "danger": ButtonStyle.danger,
    }
    return mapping.get(normalize_style_name(name), ButtonStyle.secondary)


BUTTON_TAG_RE = re.compile(r"\{\{btn:([^{}]+)\}\}")


def parse_button_tags(text: str, *, default_row: int = 0) -> tuple[str, list[PanelButton]]:
    if not text:
        return text, []

    out_buttons: list[PanelButton] = []
    base_row = max(0, min(4, default_row))

    action_aliases = {
        "order": "order",
        "ticket": "order",
        "заказ": "order",
        "support": "support",
        "help": "support",
        "саппорт": "support",
        "поддержка": "support",
        "url": "url",
        "link": "url",
        "ссылка": "url",
    }

    def _resolve_row(pos: str) -> int:
        p = (pos or "inline").strip().lower()
        if p == "bottom":
            return 4
        if p.startswith("row"):
            n = p[3:].strip()
            if n.isdigit():
                return max(0, min(4, int(n)))
        return base_row

    def _repl(match: re.Match[str]) -> str:
        raw = match.group(1).strip()
        parts = [x.strip() for x in raw.split("|")]
        if len(parts) < 2:
            return match.group(0)

        label = parts[0] or "Button"
        action_raw = parts[1].lower()
        action = action_aliases.get(action_raw, action_raw)
        style = normalize_style_name(parts[2]) if len(parts) > 2 and parts[2] else "secondary"
        pos = parts[3] if len(parts) > 3 else "inline"
        emoji = parts[4] if len(parts) > 4 else ""
        extra = parts[5] if len(parts) > 5 else ""

        row = _resolve_row(pos)

        url: str | None = None
        if action == "url":
            url = extra or None
            if not url:
                return f"{(emoji + ' ' ) if emoji else ''}{label}"

        if action not in {"order", "support", "url"}:
            return match.group(0)

        out_buttons.append(
            PanelButton(
                label=label,
                emoji=emoji,
                style=style,
                action=action,
                url=url,
                row=row,
            )
        )
        return f"{(emoji + ' ') if emoji else ''}{label}"

    cleaned = BUTTON_TAG_RE.sub(_repl, text)
    return cleaned, out_buttons


def materialize_post_for_send(post: SavedPost) -> SavedPost:
    cleaned_desc, tag_buttons = parse_button_tags(post.description, default_row=0)
    merged_buttons = list(post.panel_buttons) + tag_buttons
    return replace(post, description=cleaned_desc, panel_buttons=merged_buttons)


class TicketOpenView(discord.ui.View):
    def __init__(self, bot: "BridgeBot", post: SavedPost):
        super().__init__(timeout=None)
        self.bot = bot
        self.post = post

        buttons = list(post.panel_buttons)
        if not buttons:
            buttons = [
                PanelButton(label=post.order_label, emoji=(post.order_emoji or ""), style=post.order_style, action="order", row=0),
                PanelButton(label=post.support_label, emoji=(post.support_emoji or ""), style=post.support_style, action="support", row=0),
            ]

        buttons.sort(key=lambda b: int(b.row) if isinstance(b.row, int) else 0)
        for i, cfg in enumerate(buttons):
            action = str(cfg.action or "none").strip().lower()
            row = int(cfg.row) if isinstance(cfg.row, int) else 0
            if row < 0:
                row = 0
            if row > 4:
                row = 4

            if action == "url" and cfg.url:
                btn = discord.ui.Button(
                    label=cfg.label or "Open",
                    emoji=(cfg.emoji or None),
                    style=ButtonStyle.link,
                    url=cfg.url,
                    row=row,
                )
                self.add_item(btn)
                continue

            btn = discord.ui.Button(
                label=cfg.label or "Action",
                emoji=(cfg.emoji or None),
                style=style_from_name(cfg.style),
                custom_id=f"panel:{post.name}:{i}",
                row=row,
            )

            if action == "order":
                async def cb(interaction: discord.Interaction, ticket_type: str = "order") -> None:
                    await self._safe_create(interaction, ticket_type)
                btn.callback = cb
                self.add_item(btn)
            elif action == "support":
                async def cb(interaction: discord.Interaction, ticket_type: str = "support") -> None:
                    await self._safe_create(interaction, ticket_type)
                btn.callback = cb
                self.add_item(btn)

    async def _safe_create(self, interaction: discord.Interaction, ticket_type: str) -> None:
        try:
            await self.bot.create_ticket(interaction, ticket_type, self.post)
        except Exception as exc:
            logger.exception("Ticket button failed")
            if interaction.response.is_done():
                await interaction.followup.send(f"Ошибка открытия тикета: {exc}", ephemeral=True)
            else:
                await interaction.response.send_message(
                    f"Ошибка открытия тикета: {exc}", ephemeral=True
                )


class BridgeBot(commands.Bot):
    def __init__(self, ticket_store: TicketStore, post_store: PostStore, user_links: UserLinkStore, tg_roles: TgRoleStore):
        intents = discord.Intents.default()
        intents.message_content = True
        intents.members = True
        super().__init__(command_prefix="!", intents=intents)
        self.ticket_store = ticket_store
        self.post_store = post_store
        self.user_links = user_links
        self.tg_roles = tg_roles
        self.tg_app: Application | None = None

    async def setup_hook(self) -> None:
        self.add_view(CloseTicketButton(self))

        restored = 0
        for post in self.post_store.list_posts():
            runtime_post = materialize_post_for_send(post)
            if not (runtime_post.is_ticket_panel or runtime_post.panel_buttons):
                continue
            try:
                self.add_view(TicketOpenView(self, runtime_post))
                restored += 1
            except Exception:
                logger.exception("Failed to restore persistent panel view for post '%s'", post.name)
        if restored:
            logger.info("Restored %s persistent ticket panel view(s)", restored)

    async def send_tg(self, chat_id: int, text: str) -> None:
        if not self.tg_app:
            return
        await self.tg_app.bot.send_message(chat_id=chat_id, text=text, disable_web_page_preview=True)

    def _notification_targets(self, ticket_type: str, opener_discord_id: int) -> set[int]:
        targets: set[int] = set()

        opener_chat = self.user_links.get_tg_chat(opener_discord_id)
        if opener_chat:
            targets.add(opener_chat)

        targets.update(TG_ADMIN_IDS)
        targets.update(self.tg_roles.chats_with_roles({"admin", "manager"}))
        if ticket_type == "order":
            targets.update(self.tg_roles.chats_with_roles({"builder"}))
        return targets

    async def notify_ticket(self, ticket_type: str, opener_discord_id: int, text: str) -> None:
        targets = self._notification_targets(ticket_type, opener_discord_id)

        if not targets:
            logger.warning("No Telegram targets for ticket notification: %s", text)

        for chat_id in targets:
            try:
                await self.send_tg(chat_id, text)
            except Exception:
                logger.exception("Failed sending DM to TG user/chat %s", chat_id)

    @staticmethod
    def _message_text(msg: discord.Message) -> str:
        text = (msg.content or "").strip()
        if not text and msg.attachments:
            text = "[attachment] " + ", ".join(a.filename for a in msg.attachments)
        return text or "[empty]"

    async def _update_ticket_digest(self, channel_id: int) -> None:
        if not self.tg_app:
            return
        binding = self.ticket_store.get(channel_id)
        if not binding:
            return
        channel = self.get_channel(channel_id)
        if not isinstance(channel, discord.TextChannel):
            return

        messages = []
        async for m in channel.history(limit=20):
            if m.author.bot:
                continue
            messages.append(m)
            if len(messages) >= 5:
                break

        messages = list(reversed(messages))
        lines = [f"🧾 {binding.ticket_type.upper()} #{channel.name}"]
        if messages:
            for i, m in enumerate(messages, 1):
                lines.append(f"{i}. {m.author.display_name}: {self._message_text(m)}")
        else:
            lines.append("(пока нет сообщений клиента)")

        digest = "\n".join(lines)
        if len(digest) > 3800:
            digest = digest[:3800] + "\n..."

        targets = self._notification_targets(binding.ticket_type, binding.opener_discord_id)
        if not targets:
            return

        changed = False
        for chat_id in targets:
            key = str(chat_id)
            msg_id = binding.digest_message_ids.get(key)
            try:
                if msg_id:
                    await self.tg_app.bot.edit_message_text(
                        chat_id=chat_id,
                        message_id=msg_id,
                        text=digest,
                        disable_web_page_preview=True,
                    )
                else:
                    sent = await self.tg_app.bot.send_message(
                        chat_id=chat_id,
                        text=digest,
                        disable_web_page_preview=True,
                    )
                    binding.digest_message_ids[key] = sent.message_id
                    changed = True
            except Exception:
                logger.exception("Failed updating digest chat=%s", chat_id)

        if changed:
            self.ticket_store.set(binding)

    async def _collect_ticket_archive(self, channel: discord.TextChannel) -> list[Path]:
        ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        ticket_dir = ARCHIVE_DIR / f"ticket_{channel.id}_{ts}"
        attach_dir = ticket_dir / "attachments"
        attach_dir.mkdir(parents=True, exist_ok=True)

        lines: list[str] = []
        json_messages: list[dict] = []
        files: list[Path] = []
        downloaded: list[Path] = []

        async with ClientSession() as session:
            async for m in channel.history(limit=None, oldest_first=True):
                t = m.created_at.astimezone(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
                msg_line = f"[{t}] {m.author.display_name} ({m.author.id}): {m.content or ''}".rstrip()
                lines.append(msg_line)

                embeds_payload = []
                for emb in m.embeds:
                    try:
                        embeds_payload.append(emb.to_dict())
                    except Exception:
                        embeds_payload.append({"error": "embed_to_dict_failed"})

                json_messages.append({
                    "id": m.id,
                    "created_at": t,
                    "author": {
                        "id": m.author.id,
                        "display_name": m.author.display_name,
                    },
                    "content": m.content or "",
                    "attachments": [
                        {
                            "filename": a.filename,
                            "url": a.url,
                            "content_type": a.content_type,
                            "size": a.size,
                        }
                        for a in m.attachments
                    ],
                    "embeds": embeds_payload,
                })

                for a in m.attachments:
                    fname = sanitize_filename(a.filename)
                    ap = attach_dir / f"{m.id}_{fname}"
                    try:
                        async with session.get(a.url) as resp:
                            if resp.status == 200:
                                ap.write_bytes(await resp.read())
                                downloaded.append(ap)
                                lines.append(f"  [file] {fname} -> {ap.name}")
                            else:
                                lines.append(f"  [file] {fname} -> download failed status={resp.status}")
                    except Exception as exc:
                        logger.exception("Attachment download failed: %s", a.url)
                        lines.append(f"  [file] {fname} -> download error: {exc}")

        txt_path = ticket_dir / "dialog.txt"
        txt_path.write_text("\n".join(lines), encoding="utf-8")
        files.append(txt_path)

        json_path = ticket_dir / "dialog.json"
        json_path.write_text(json.dumps(json_messages, ensure_ascii=False, indent=2), encoding="utf-8")
        files.append(json_path)

        try:
            from docx import Document

            doc = Document()
            doc.add_heading(f"Ticket transcript #{channel.id}", 0)
            for line in lines:
                doc.add_paragraph(line)
            docx_path = ticket_dir / "dialog.docx"
            doc.save(docx_path)
            files.append(docx_path)
        except Exception:
            logger.exception("DOCX export failed")

        files.extend(downloaded)

        zip_path = ticket_dir / "ticket_archive.zip"
        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for fp in files:
                zf.write(fp, arcname=fp.relative_to(ticket_dir))
        files.append(zip_path)
        return files

    async def _send_ticket_rating(self, chat_id: int, ticket_type: str, channel_name: str, channel_id: int) -> None:
        if not self.tg_app:
            return
        kb = InlineKeyboardMarkup(
            [[
                InlineKeyboardButton("✅ Заказ успешно", callback_data=f"ticket_rate:{channel_id}:success"),
                InlineKeyboardButton("➖ Нейтрально", callback_data=f"ticket_rate:{channel_id}:neutral"),
                InlineKeyboardButton("❌ Не выполнен", callback_data=f"ticket_rate:{channel_id}:failed"),
            ]]
        )
        await self.tg_app.bot.send_message(
            chat_id=chat_id,
            text=f"Тикет закрыт: {ticket_type.upper()} #{channel_name}\nОцените результат:",
            reply_markup=kb,
        )

    async def close_ticket(self, channel: discord.abc.GuildChannel, binding: TicketBinding, closed_by: discord.abc.User) -> None:
        if not isinstance(channel, discord.TextChannel):
            return

        await self.notify_ticket(
            ticket_type=binding.ticket_type,
            opener_discord_id=binding.opener_discord_id,
            text=f"🔒 Тикет закрыт: #{channel.name} (закрыл: {closed_by.display_name})",
        )

        archive_files = await self._collect_ticket_archive(channel)
        targets = self._notification_targets(binding.ticket_type, binding.opener_discord_id)

        for chat_id in targets:
            for fp in archive_files:
                try:
                    with fp.open("rb") as f:
                        await self.tg_app.bot.send_document(chat_id=chat_id, document=f)
                except Exception:
                    logger.exception("Failed sending archive file %s to chat %s", fp, chat_id)
            try:
                await self._send_ticket_rating(chat_id, binding.ticket_type, channel.name, channel.id)
            except Exception:
                logger.exception("Failed sending rating keyboard to chat %s", chat_id)

        self.ticket_store.remove(channel.id)
        await channel.delete(reason="Ticket closed")

    async def create_ticket(self, interaction: discord.Interaction, ticket_type: str, post: SavedPost | None = None) -> None:
        if not interaction.guild or not interaction.user or not interaction.channel:
            return
        await interaction.response.defer(ephemeral=True)

        member = interaction.user
        base_name = re.sub(r"[^a-zA-Z0-9\-_]", "-", member.display_name).strip("-").lower() or "user"
        prefix = ORDER_PREFIX if ticket_type == "order" else SUPPORT_PREFIX
        channel_name = f"{prefix}-{base_name}"[:95]

        category = interaction.guild.get_channel(TICKET_CATEGORY_ID) if TICKET_CATEGORY_ID else interaction.channel.category
        support_role = interaction.guild.get_role(SUPPORT_ROLE_ID) if SUPPORT_ROLE_ID else None
        me = interaction.guild.me
        if not me:
            await interaction.followup.send("Не найден bot member", ephemeral=True)
            return

        overwrites: dict[discord.Role | discord.Member, discord.PermissionOverwrite] = {
            interaction.guild.default_role: discord.PermissionOverwrite(view_channel=False),
            member: discord.PermissionOverwrite(view_channel=True, send_messages=True, read_message_history=True),
            me: discord.PermissionOverwrite(view_channel=True, send_messages=True, manage_channels=True),
        }
        if support_role:
            overwrites[support_role] = discord.PermissionOverwrite(view_channel=True, send_messages=True, read_message_history=True)

        ticket_channel = await interaction.guild.create_text_channel(
            name=channel_name,
            category=category if isinstance(category, discord.CategoryChannel) else None,
            overwrites=overwrites,
            topic=f"{ticket_type} ticket opened by {member} ({member.id})",
            reason="New ticket",
        )

        self.ticket_store.set(TicketBinding(ticket_channel_id=ticket_channel.id, opener_discord_id=member.id, ticket_type=ticket_type, digest_message_ids={}))

        auto_msg = ""
        if post:
            auto_msg = post.auto_order_message if ticket_type == "order" else post.auto_support_message

        if auto_msg:
            description = auto_msg.replace("{user}", member.mention)
        else:
            description = f"Привет, {member.mention}! Опишите задачу подробно."

        embed = discord.Embed(
            title=f"{ticket_type.upper()} ticket",
            description=description,
            color=0x5865F2,
        )
        await ticket_channel.send(content=member.mention, embed=embed, view=CloseTicketButton(self))

        await self.notify_ticket(
            ticket_type=ticket_type,
            opener_discord_id=member.id,
            text=(
                f"🆕 Тикет открыт: {ticket_type.upper()}\n"
                f"Пользователь: {member.display_name} ({member.id})\n"
                f"Discord канал: #{ticket_channel.name} ({ticket_channel.id})"
            ),
        )
        await interaction.followup.send(f"Тикет создан: {ticket_channel.mention}", ephemeral=True)



    async def on_message(self, message: discord.Message) -> None:
        if message.author.bot:
            return
        await self._update_ticket_digest(message.channel.id)
        await self.process_commands(message)

    async def on_message_edit(self, before: discord.Message, after: discord.Message) -> None:
        if after.author and after.author.bot:
            return
        await self._update_ticket_digest(after.channel.id)

    async def on_message_delete(self, message: discord.Message) -> None:
        await self._update_ticket_digest(message.channel.id)

    async def on_raw_message_delete(self, payload: discord.RawMessageDeleteEvent) -> None:
        await self._update_ticket_digest(payload.channel_id)


def hex_midpoint(a: str, b: str) -> str:
    a = a.strip("#")
    b = b.strip("#")
    ar, ag, ab = int(a[0:2], 16), int(a[2:4], 16), int(a[4:6], 16)
    br, bg, bb = int(b[0:2], 16), int(b[2:4], 16), int(b[4:6], 16)
    return f"{(ar+br)//2:02X}{(ag+bg)//2:02X}{(ab+bb)//2:02X}"


def embeds_from_post(post: SavedPost) -> list[discord.Embed]:
    base_color_hex = post.color_hex
    if post.auto_gradient:
        base_color_hex = hex_midpoint(post.gradient_start_hex, post.gradient_end_hex)
    main = discord.Embed(
        title=post.title,
        description=post.description,
        color=int(base_color_hex.strip("#"), 16),
    )
    if post.image_url and post.image_position == "top":
        main.description = f"[image above]\n{main.description}"
        main.set_image(url=post.image_url)
    elif post.image_url:
        main.set_image(url=post.image_url)

    result = [main]
    for b in post.extra_blocks:
        eb = discord.Embed(
            title=b.title,
            description=b.description,
            color=int(b.color_hex.strip("#"), 16),
        )
        if b.image_url:
            if b.image_position == "top":
                eb.description = f"[image above]\n{eb.description}"
            eb.set_image(url=b.image_url)
        result.append(eb)
    return result


async def publish_post(bot: BridgeBot, post: SavedPost) -> discord.Message:
    channel = bot.get_channel(post.channel_id)
    if not isinstance(channel, discord.TextChannel):
        raise ValueError(f"Канал {post.channel_id} не найден")

    runtime_post = materialize_post_for_send(post)
    embeds = embeds_from_post(runtime_post)
    has_runtime_buttons = bool(runtime_post.panel_buttons)
    if runtime_post.is_ticket_panel or has_runtime_buttons:
        view = TicketOpenView(bot, runtime_post)
        return await channel.send(embeds=embeds, view=view)
    return await channel.send(embeds=embeds)


async def tg_reply(update: Update, text: str) -> None:
    msg = update.effective_message
    if msg:
        await msg.reply_text(text)


def can_manage_roles(update: Update) -> bool:
    uid = update.effective_user.id if update.effective_user else None
    return bool(uid and uid in TG_ADMIN_IDS)


async def tg_start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await tg_reply(
        update,
        "Команды:\n"
        "/bind_discord <discord_user_id>\n"
        "/set_role <tg_user_id> <admin|manager|builder|viewer> [chat_id]\n"
        "/register_me <role> - зарегистрировать себя и чат для уведомлений\n"
        "/my_role\n"
        "/post_save <name> <channel_id> title|description|color_hex|image_url?\n"
        "/post_send <name>\n"
        "/reply_ticket <discord_channel_id> <text>",
    )


async def tg_bind_discord(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    bot: BridgeBot = context.application.bot_data["discord_bot"]
    if not update.effective_chat or len(context.args) != 1:
        await tg_reply(update, "Пример: /bind_discord 123456789012345678")
        return
    try:
        discord_user_id = int(context.args[0])
        bot.user_links.link(discord_user_id, update.effective_chat.id)
        if update.effective_user:
            bot.tg_roles.register_chat(update.effective_user.id, update.effective_chat.id)
        await tg_reply(update, f"Привязано ✅ discord={discord_user_id} -> tg_chat={update.effective_chat.id}")
    except Exception as exc:
        await tg_reply(update, f"Ошибка: {exc}")


async def tg_set_role(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    bot: BridgeBot = context.application.bot_data["discord_bot"]
    if not can_manage_roles(update):
        await tg_reply(update, "Нет прав. Добавь свой TG user id в TG_ADMIN_IDS")
        return
    if len(context.args) not in {2, 3}:
        await tg_reply(update, "Пример: /set_role 123456789 manager [chat_id]")
        return
    tg_user_id = int(context.args[0])
    role = context.args[1].lower()
    chat_id = int(context.args[2]) if len(context.args) == 3 else None
    if role not in {"admin", "manager", "builder", "viewer"}:
        await tg_reply(update, "Роль: admin|manager|builder|viewer")
        return
    bot.tg_roles.set_role(tg_user_id, role, chat_id=chat_id)
    msg = f"OK: {tg_user_id} -> {role}" + (f", chat_id={chat_id}" if chat_id is not None else "")
    await tg_reply(update, msg)


async def tg_my_role(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    bot: BridgeBot = context.application.bot_data["discord_bot"]
    uid = update.effective_user.id if update.effective_user else 0
    await tg_reply(update, f"Твоя роль: {bot.tg_roles.get_role(uid)}")

async def tg_register_me(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    bot: BridgeBot = context.application.bot_data["discord_bot"]
    if not update.effective_user or not update.effective_chat:
        return
    if len(context.args) != 1:
        await tg_reply(update, "Пример: /register_me builder")
        return

    role = context.args[0].lower().strip()
    if role not in {"admin", "manager", "builder", "viewer"}:
        await tg_reply(update, "Роль: admin|manager|builder|viewer")
        return

    # regular users can self-register only viewer/builder; admin/manager only via /set_role
    if role in {"admin", "manager"} and not can_manage_roles(update):
        await tg_reply(update, "Эту роль выдает только админ")
        return

    bot.tg_roles.set_role(update.effective_user.id, role, update.effective_chat.id)
    await tg_reply(update, f"Зарегистрирован ✅ role={role}, chat_id={update.effective_chat.id}")




def parse_post_args(raw: str) -> tuple[str, str, str, str | None]:
    parts = [p.strip() for p in raw.split("|")]
    if len(parts) < 3:
        raise ValueError("Нужно минимум title|description|color")
    title, desc, color = parts[0], parts[1], parts[2]
    image = parts[3] if len(parts) > 3 and parts[3] else None
    return title, desc, color, image


async def tg_post_save(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    bot: BridgeBot = context.application.bot_data["discord_bot"]
    if len(context.args) < 3:
        await tg_reply(update, "Пример: /post_save name 123 title|desc|2ECC71|https://img")
        return
    try:
        name = context.args[0].strip().lower()
        channel_id = int(context.args[1])
        title, desc, color, image = parse_post_args(" ".join(context.args[2:]))
        post = SavedPost(name=name, channel_id=channel_id, title=title, description=desc, color_hex=color, image_url=image)
        msg = await publish_post(bot, post)
        post.last_message_id = msg.id
        bot.post_store.set(post)
        await tg_reply(update, f"Сохранено и отправлено ✅\nchannel_id={channel_id}\nmessage_id={msg.id}")
    except Exception as exc:
        await tg_reply(update, f"Ошибка: {exc}")


async def tg_post_send(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    bot: BridgeBot = context.application.bot_data["discord_bot"]
    if len(context.args) != 1:
        await tg_reply(update, "Пример: /post_send name")
        return
    post = bot.post_store.get(context.args[0].strip().lower())
    if not post:
        await tg_reply(update, "Шаблон не найден")
        return
    try:
        msg = await publish_post(bot, post)
        post.last_message_id = msg.id
        bot.post_store.set(post)
        await tg_reply(update, f"Опубликовано ✅\nchannel_id={post.channel_id}\nmessage_id={msg.id}")
    except Exception as exc:
        await tg_reply(update, f"Ошибка: {exc}")


async def tg_reply_ticket(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    bot: BridgeBot = context.application.bot_data["discord_bot"]
    role = bot.tg_roles.get_role(update.effective_user.id if update.effective_user else 0)
    if role not in {"admin", "manager", "builder"}:
        await tg_reply(update, "Нет прав на reply_ticket")
        return
    if len(context.args) < 2:
        await tg_reply(update, "Пример: /reply_ticket 123456789012345678 текст")
        return
    try:
        channel_id = int(context.args[0])
        text = " ".join(context.args[1:]).strip()
        ch = bot.get_channel(channel_id)
        if not isinstance(ch, discord.TextChannel):
            await tg_reply(update, "Канал не найден")
            return
        uname = update.effective_user.full_name if update.effective_user else "TG user"
        await ch.send(f"📨 TG {uname}: {text}")
        await tg_reply(update, "Отправлено ✅")
    except Exception as exc:
        await tg_reply(update, f"Ошибка: {exc}")




async def tg_ticket_rate(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    query = update.callback_query
    if not query or not query.data:
        return

    await query.answer()
    parts = query.data.split(":", 2)
    if len(parts) != 3:
        return

    status_map = {
        "success": "✅ Заказ успешно",
        "neutral": "➖ Нейтрально",
        "failed": "❌ Не выполнен",
    }
    status = status_map.get(parts[2], parts[2])

    text = query.message.text if query.message and query.message.text else "Тикет оценен"
    new_text = f"{text}\n\nВыбрано: {status}"
    try:
        await query.edit_message_text(new_text)
    except Exception:
        # fallback: hide buttons at least
        await query.edit_message_reply_markup(reply_markup=None)

async def tg_error_handler(update: object, context: ContextTypes.DEFAULT_TYPE) -> None:
    logger.exception("Telegram error", exc_info=context.error)
    if isinstance(update, Update):
        await tg_reply(update, f"Ошибка: {context.error}")



def parse_post_json(data: dict) -> SavedPost:
    blocks = [ExtraBlock(**b) for b in data.get("extra_blocks", []) if isinstance(b, dict)]
    panel_buttons: list[PanelButton] = []
    for b in data.get("panel_buttons", []) or []:
        if not isinstance(b, dict):
            continue
        try:
            row = int(b.get("row", 0))
        except Exception:
            row = 0
        panel_buttons.append(
            PanelButton(
                label=str(b.get("label", "Button")),
                emoji=str(b.get("emoji", "")).strip(),
                style=normalize_style_name(str(b.get("style", "secondary"))),
                action=str(b.get("action", "none")).strip().lower(),
                url=str(b.get("url", "")).strip() or None,
                row=max(0, min(4, row)),
            )
        )

    post = SavedPost(
        name=str(data["name"]).strip().lower(),
        channel_id=int(data["channel_id"]),
        title=str(data.get("title", "")),
        description=str(data.get("description", "")),
        color_hex=str(data.get("color_hex", "2ECC71")),
        image_url=str(data.get("image_url", "")).strip() or None,
        image_position=str(data.get("image_position", "bottom")),
        is_ticket_panel=bool(data.get("is_ticket_panel", False)),
        order_label=str(data.get("order_label", "ORDER")),
        order_emoji=str(data.get("order_emoji", "🧾")).strip(),
        order_style=normalize_style_name(str(data.get("order_style", "success"))),
        support_label=str(data.get("support_label", "SUPPORT")),
        support_emoji=str(data.get("support_emoji", "🛟")).strip(),
        support_style=normalize_style_name(str(data.get("support_style", "primary"))),
        auto_gradient=bool(data.get("auto_gradient", False)),
        gradient_start_hex=str(data.get("gradient_start_hex", "2ECC71")),
        gradient_end_hex=str(data.get("gradient_end_hex", "5865F2")),
        auto_order_message=str(data.get("auto_order_message", "")),
        auto_support_message=str(data.get("auto_support_message", "")),
        panel_buttons=panel_buttons,
        extra_blocks=blocks,
        last_message_id=int(data["last_message_id"]) if data.get("last_message_id") else None,
    )
    int(post.color_hex.strip("#"), 16)
    int(post.gradient_start_hex.strip("#"), 16)
    int(post.gradient_end_hex.strip("#"), 16)
    return post


async def web_index(request: web.Request) -> web.Response:
    return web.Response(text=EDITOR_HTML_PATH.read_text(encoding="utf-8"), content_type="text/html")


async def web_list_posts(request: web.Request) -> web.Response:
    bot: BridgeBot = request.app["bot"]
    return web.json_response([asdict(x) for x in bot.post_store.list_posts()])


async def web_get_post(request: web.Request) -> web.Response:
    bot: BridgeBot = request.app["bot"]
    post = bot.post_store.get(request.match_info["name"])
    if not post:
        return web.json_response({"error": "not_found"}, status=404)
    return web.json_response(asdict(post))


async def web_save_post(request: web.Request) -> web.Response:
    bot: BridgeBot = request.app["bot"]
    try:
        post = parse_post_json(await request.json())
        bot.post_store.set(post)
        return web.json_response({"ok": True, "post": asdict(post)})
    except Exception as exc:
        return web.json_response({"error": str(exc)}, status=400)


async def web_publish_post(request: web.Request) -> web.Response:
    bot: BridgeBot = request.app["bot"]
    post = bot.post_store.get(request.match_info["name"])
    if not post:
        return web.json_response({"error": "not_found"}, status=404)
    try:
        msg = await publish_post(bot, post)
        post.last_message_id = msg.id
        bot.post_store.set(post)
        return web.json_response({"ok": True, "channel_id": post.channel_id, "message_id": msg.id})
    except Exception as exc:
        return web.json_response({"error": str(exc)}, status=400)


async def start_web_server(bot: BridgeBot) -> web.AppRunner:
    app = web.Application()
    app["bot"] = bot
    app.router.add_get("/", web_index)
    app.router.add_get("/api/posts", web_list_posts)
    app.router.add_get("/api/posts/{name}", web_get_post)
    app.router.add_post("/api/posts/save", web_save_post)
    app.router.add_post("/api/posts/{name}/publish", web_publish_post)

    runner = web.AppRunner(app)
    await runner.setup()
    await web.TCPSite(runner, WEB_HOST, WEB_PORT).start()
    logger.info("Web editor started on http://%s:%s", WEB_HOST, WEB_PORT)
    return runner


def validate_env() -> None:
    if not DISCORD_TOKEN or not TELEGRAM_TOKEN:
        raise RuntimeError("Set DISCORD_TOKEN and TELEGRAM_TOKEN in .env")


async def run() -> None:
    validate_env()
    ticket_store = TicketStore(STATE_FILE)
    post_store = PostStore(POSTS_FILE)
    user_links = UserLinkStore(USER_LINKS_FILE)
    tg_roles = TgRoleStore(TG_ROLES_FILE)
    d_bot = BridgeBot(ticket_store, post_store, user_links, tg_roles)

    tg_app = Application.builder().token(TELEGRAM_TOKEN).build()
    d_bot.tg_app = tg_app
    tg_app.bot_data["discord_bot"] = d_bot

    tg_app.add_handler(CommandHandler("start", tg_start))
    tg_app.add_handler(CommandHandler("bind_discord", tg_bind_discord))
    tg_app.add_handler(CommandHandler("set_role", tg_set_role))
    tg_app.add_handler(CommandHandler("register_me", tg_register_me))
    tg_app.add_handler(CommandHandler("my_role", tg_my_role))
    tg_app.add_handler(CommandHandler("post_save", tg_post_save))
    tg_app.add_handler(CommandHandler("post_send", tg_post_send))
    tg_app.add_handler(CommandHandler("reply_ticket", tg_reply_ticket))
    tg_app.add_handler(CallbackQueryHandler(tg_ticket_rate, pattern=r"^ticket_rate:"))
    tg_app.add_error_handler(tg_error_handler)

    await tg_app.bot.set_my_commands(
        [
            BotCommand("start", "Показать команды"),
            BotCommand("bind_discord", "Привязать Discord пользователя"),
            BotCommand("set_role", "Назначить роль пользователю TG"),
            BotCommand("register_me", "Зарегистрировать себя для уведомлений"),
            BotCommand("my_role", "Показать мою роль"),
            BotCommand("post_save", "Сохранить и отправить шаблон"),
            BotCommand("post_send", "Опубликовать сохраненный шаблон"),
            BotCommand("reply_ticket", "Отправить сообщение в Discord тикет"),
        ]
    )

    await tg_app.initialize()
    await tg_app.start()
    await tg_app.updater.start_polling()

    web_runner = await start_web_server(d_bot)
    discord_task = asyncio.create_task(d_bot.start(DISCORD_TOKEN))

    logger.info("Discord + Telegram + Web editor started")
    try:
        await discord_task
    finally:
        await web_runner.cleanup()
        await tg_app.updater.stop()
        await tg_app.stop()
        await tg_app.shutdown()


if __name__ == "__main__":
    try:
        asyncio.run(run())
    except Exception:
        logger.exception("Fatal startup/runtime error")
        raise
